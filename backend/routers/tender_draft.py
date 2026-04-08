from fastapi import APIRouter, HTTPException, Depends, Body, Request, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional, List, Any
import io
import datetime
import json
import os
import uuid
import logging
from pathlib import Path

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import pdfplumber
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Import AI SDKs
import anthropic
import google.generativeai as genai
from groq import Groq

router = APIRouter()
logger = logging.getLogger(__name__)

# Configure AI Clients
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

claude_client = None
if ANTHROPIC_API_KEY and "ENTER_YOUR" not in ANTHROPIC_API_KEY:
    try:
        claude_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    except Exception as e:
        logger.error(f"Failed to initialize Anthropic client: {e}")

gemini_model = None
if GEMINI_API_KEY and "ENTER_YOUR" not in GEMINI_API_KEY:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        gemini_model = genai.GenerativeModel('gemini-1.5-pro')
    except Exception as e:
        logger.error(f"Failed to initialize Gemini client: {e}")

groq_client = None
if GROQ_API_KEY and "ENTER_YOUR" not in GROQ_API_KEY:
    try:
        groq_client = Groq(api_key=GROQ_API_KEY)
    except Exception as e:
        logger.error(f"Failed to initialize Groq client: {e}")

# --- HELPER FUNCTIONS FROM PROVIDED CODE ---

def register_regional_font():
    # Windows system font paths - prioritizing Gautami for Telugu as per RTGS requirements
    search_paths = [
        os.path.join(os.environ.get('WINDIR', 'C:\\Windows'), 'Fonts', f)
        for f in ['gautami.ttf', 'nirmala.ttc', 'nirmala.ttf', 'mangal.ttf', 'arialuni.ttf']
    ]
    for path in search_paths:
        if os.path.exists(path):
            try:
                font_name = Path(path).stem
                pdfmetrics.registerFont(TTFont(font_name, path))
                logger.info(f"Registered regional font: {font_name}")
                return font_name
            except Exception as e:
                logger.error(f"Failed to register font {path}: {e}")
                continue
    return 'Helvetica'

REGIONAL_FONT = register_regional_font()

def find_logo_path():
    this_file = Path(__file__).resolve()
    # Search in various potential locations
    paths = [
        this_file.parent.parent / "uploads" / "images.jpg",
        Path(os.getcwd()) / "uploads" / "images.jpg",
        Path(os.getcwd()) / "backend" / "uploads" / "images.jpg",
        this_file.parent.parent.parent / "uploads" / "images.jpg",
    ]
    for p in paths:
        if p.exists():
            return str(p)
    return None

def extract_pdf_text(pdf_path):
    try:
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            # Limit to first 15 pages to keep prompt size reasonable
            for page in pdf.pages[:15]:
                text += (page.extract_text() or "") + "\n"
        return text
    except Exception as e:
        logger.error(f"PDF Extraction Error: {e}")
        return ""

# --- MODELS ---

class EnquiryForm(BaseModel):
    stage: str
    name: str
    type: str
    mandatory: str

class TenderExtractedData(BaseModel):
    bidNumber: str
    dated: str
    endDate: str
    openingDate: str
    validity: str
    ministry: str
    department: str
    organisation: str
    office: str
    category: str
    estimatedValue: str
    emdRequired: str
    advisoryBank: str
    epbgPercentage: str
    epbgDuration: str
    beneficiary: str
    miiPreference: str
    msePreference: str
    bidderTurnover: str
    oemTurnover: str
    experience: str
    mseRelaxation: str
    startupRelaxation: str
    pastPerformance: str
    docRequired: str
    techClarificationTime: str
    minBidsForAutoExtend: str
    autoExtendDays: str
    autoExtensionCount: str
    priceBreakupRequired: str
    transactionFee: str
    plCode: Optional[str] = None
    approvingAgency: Optional[str] = None
    uvamItemId: Optional[str] = None
    enquiryForms: List[EnquiryForm] = []
    clauses: str
    source: str
    source_label: str

class UnifiedDownloadRequest(TenderExtractedData):
    format: str = "pdf"

# --- PDF GENERATORS ---

def generate_ap_eproc_pdf_buffer(data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=20, leftMargin=20, topMargin=20, bottomMargin=20)
    styles = getSampleStyleSheet()
    
    SECTION_HEADER_BG = colors.HexColor('#e8f0fe') 
    SECTION_HEADER_TEXT = colors.HexColor('#1a73e8') 
    LABEL_TEXT_COLOR = colors.HexColor('#444444')
    BORDER_COLOR = colors.HexColor('#e0e0e0')

    title_style = ParagraphStyle('Title', fontName='Helvetica-Bold', fontSize=16, textColor=colors.HexColor('#0056b3'), alignment=1, spaceAfter=2)
    section_style = ParagraphStyle('Section', fontSize=10, textColor=SECTION_HEADER_TEXT, fontName='Helvetica-Bold')
    label_style = ParagraphStyle('Label', fontSize=8, textColor=LABEL_TEXT_COLOR, fontName='Helvetica-Bold', leading=10)
    value_style = ParagraphStyle('Value', fontSize=8, fontName='Helvetica', leading=10)
    
    elements = []
    
    logo_path = find_logo_path()
    if logo_path:
        elements.append(Image(logo_path, width=120, height=45))
        elements.append(Spacer(1, 10))

    elements.append(Paragraph("eGP Portal - Tenders", title_style))
    elements.append(Paragraph("Government of Andhra Pradesh", ParagraphStyle('Sub', alignment=1, fontSize=12, spaceAfter=20)))

    sections = [
        ('Current Tender Details', {
            'Tender ID': data.get('bidNumber'),
            'Tender Notice Number': f"NIT/{data.get('bidNumber')}",
            'Name of Work': data.get('category'),
            'Tender Category': 'WORKS',
            'Tender Type': 'OPEN - NCB',
            'Estimated Contract Value': f"₹ {data.get('estimatedValue')}",
            'Submission Closing Date': data.get('endDate'),
            'Evaluation Type': data.get('evaluationMethod')
        }),
        ('Enquiry Particulars', {
            'Department Name': data.get('department'),
            'Circle/Division': data.get('organisation'),
            'Name of Project': data.get('projectName', 'As per NIT'),
            'Period of Completion': f"{data.get('epbgDuration')} Months",
            'Bidding Type': data.get('bidType', 'Two Packet'),
            'Currency Type': 'INR',
            'Consortium/ JV': 'Not Applicable'
        }),
        ('Tender Dates', {
            'Start Date & Time': data.get('dated'),
            'End Date & Time': data.get('endDate'),
            'Closing Date & Time': data.get('endDate'),
            'Bid Validity Period': data.get('validity'),
            'Display Rank': 'Lowest'
        }),
        ('Bid Security Details', {
            'Bid Security (INR)': data.get('emdRequired'),
            'Bid Security In Favour Of': data.get('beneficiary'),
            'Mode of Payment': 'Online / BG'
        }),
        ('Transaction Fee Details', {
            'Transaction Fee Payable to APTS': data.get('transactionFee', '29500 (INR)')
        })
    ]

    for name, content in sections:
        elements.append(Spacer(1, 10))
        table_rows = [[Paragraph(name, section_style), '', '', '']]
        
        items = list(content.items())
        for i in range(0, len(items), 2):
            item1 = items[i]
            item2 = items[i+1] if i+1 < len(items) else (None, None)
            
            row = [Paragraph(item1[0], label_style), Paragraph(str(item1[1]), value_style)]
            row += [Paragraph(item2[0], label_style), Paragraph(str(item2[1]), value_style)] if item2[0] else ['', '']
            table_rows.append(row)

        avail_width = A4[0] - 40
        t = Table(table_rows, colWidths=[avail_width*0.2, avail_width*0.3, avail_width*0.2, avail_width*0.3])
        t.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, BORDER_COLOR),
            ('SPAN', (0,0), (3,0)),
            ('BACKGROUND', (0,0), (-1,0), SECTION_HEADER_BG),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('TOPPADDING', (0,0), (-1,-1), 6),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ]))
        elements.append(t)

    forms = data.get('enquiryForms', [])
    if forms:
        elements.append(Spacer(1, 10))
        form_rows = [[Paragraph("Enquiry Forms", section_style), '', '', '']]
        form_rows.append([Paragraph("<b>Stage</b>", label_style), Paragraph("<b>Form Name</b>", label_style), Paragraph("<b>Type</b>", label_style), Paragraph("<b>Mandatory</b>", label_style)])
        for f in forms:
            form_rows.append([
                Paragraph(f.get('stage', 'PQ'), value_style),
                Paragraph(f.get('name', 'N/A'), value_style),
                Paragraph(f.get('type', 'Standard'), value_style),
                Paragraph(f.get('mandatory', 'No'), value_style)
            ])
        t_forms = Table(form_rows, colWidths=[avail_width*0.2, avail_width*0.4, avail_width*0.2, avail_width*0.2])
        t_forms.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, BORDER_COLOR),
            ('SPAN', (0,0), (3,0)),
            ('BACKGROUND', (0,0), (-1,0), SECTION_HEADER_BG),
        ]))
        elements.append(t_forms)

    elements.append(Spacer(1, 10))
    doc_header = [[Paragraph("Required Tender Documents Details", section_style), '', '', '']]
    doc_header.append([Paragraph("<b>S.No</b>", label_style), Paragraph("<b>Document Name</b>", label_style), Paragraph("<b>Stage</b>", label_style), Paragraph("<b>Document Type</b>", label_style)])
    doc_header.append([Paragraph("1", value_style), Paragraph(data.get('docRequired', 'As per NIT'), value_style), Paragraph("COMMON", value_style), Paragraph("Mandatory", value_style)])
    
    t_doc = Table(doc_header, colWidths=[avail_width*0.1, avail_width*0.5, avail_width*0.2, avail_width*0.2])
    t_doc.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, BORDER_COLOR),
        ('SPAN', (0,0), (3,0)),
        ('BACKGROUND', (0,0), (-1,0), SECTION_HEADER_BG),
    ]))
    elements.append(t_doc)

    elements.append(Spacer(1, 40))
    elements.append(Paragraph("--- Generated via AI Procurement System ---", ParagraphStyle('Footer', alignment=1, fontSize=8, textColor=colors.grey)))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_ireps_pdf_buffer(data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=20, leftMargin=20, topMargin=20, bottomMargin=30)
    
    header_style = ParagraphStyle('IREPS_Header', fontSize=12, fontName='Helvetica-Bold', alignment=1, spaceAfter=2)
    sub_header_style = ParagraphStyle('IREPS_SubHeader', fontSize=10, fontName='Helvetica-Bold', alignment=1, spaceAfter=10)
    label_style = ParagraphStyle('IREPS_Label', fontSize=8, fontName='Helvetica-Bold', leading=9)
    value_style = ParagraphStyle('IREPS_Value', fontSize=8, fontName='Helvetica', leading=9)

    elements = []
    
    railway_name = data.get('organisation', 'NORTH WESTERN RAILWAY').upper()
    elements.append(Paragraph(f"STORES/{railway_name}", header_style))
    elements.append(Paragraph("TENDER DOCUMENT", sub_header_style))
    
    bid_info_table = [[
        Paragraph(f"<b>Tender No:</b> {data.get('bidNumber')}", value_style),
        Paragraph(f"<b>Closing Date/Time:</b> {data.get('endDate')}", value_style)
    ]]
    t_bid = Table(bid_info_table, colWidths=[250, 300])
    t_bid.setStyle(TableStyle([('ALIGN', (1,0), (1,0), 'RIGHT')]))
    elements.append(t_bid)
    elements.append(Spacer(1, 10))

    elements.append(Paragraph("1. NIT HEADER", label_style))
    nit_data = [
        ["Bidding type", "Normal Tender", "Template", "Normal"],
        ["Contract type", "Goods", "Contract Category", "Expenditure"],
        ["Tender No", data.get('bidNumber'), "Tender Type", "Open - Indigenous"],
        ["Evaluation Criteria", data.get('evaluationMethod', 'L1'), "Bidding System", "Single Packet"],
        ["Pre-Bid Required", "No", "Inspection Agency", "CONSIGNEE"],
        ["Closing Date Time", data.get('endDate'), "Ranking Order", "Lowest to Highest"],
        ["Tender Doc Cost", "0.00", "Earnest Money", data.get('emdRequired')],
        ["Tender Title", data.get('category'), "", ""]
    ]
    nit_rows = []
    for row in nit_data:
        nit_rows.append([Paragraph(f"<b>{row[0]}</b>", label_style), Paragraph(str(row[1]), value_style), Paragraph(f"<b>{row[2]}</b>", label_style), Paragraph(str(row[3]), value_style)])
    
    t_nit = Table(nit_rows, colWidths=[100, 175, 100, 175])
    t_nit.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, colors.black), ('VALIGN', (0,0), (-1,-1), 'TOP')]))
    elements.append(t_nit)
    elements.append(Spacer(1, 10))

    elements.append(Paragraph("2. ITEM DETAILS", label_style))
    item_header = [["S.No", "PL Code", "Item Type", "Stock/Non-Stock", "Approving Agency", "Qty"]]
    item_row = [
        "1", 
        data.get('plCode', '11377045'), 
        "Goods", 
        "Stock", 
        data.get('approvingAgency', 'RDSO'), 
        data.get('quantity', '60')
    ]
    t_item = Table(item_header + [item_row], colWidths=[40, 80, 80, 100, 100, 50])
    t_item.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, colors.black), ('BACKGROUND', (0,0), (-1,0), colors.lightgrey)]))
    elements.append(t_item)
    
    elements.append(Spacer(1, 5))
    elements.append(Paragraph(f"<b>Description:</b> {data.get('category')}", value_style))
    elements.append(Paragraph(f"<b>Consignee:</b> {data.get('office', 'Authority Office')}, {data.get('organisation')}", value_style))
    elements.append(Spacer(1, 10))

    for title, key in [("4. ELIGIBILITY CONDITIONS", "docRequired"), ("5. COMPLIANCE CONDITIONS", "clauses")]:
        elements.append(Paragraph(title, label_style))
        content = data.get(key, 'As per IRS Conditions')
        t_comp = Table([[Paragraph(content.replace('\n', '<br/>'), value_style)]], colWidths=[550])
        t_comp.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, colors.black)]))
        elements.append(t_comp)
        elements.append(Spacer(1, 10))

    elements.append(Paragraph("6. Bidder Certifications", label_style))
    certs = [
        "I/we the tenderer (s) am/are signing this document after carefully reading the contents.",
        "I/We also accept all the conditions of the tender and have signed all the pages in confirmation thereof.",
        "I/we declare that I/we have downloaded the tender documents from Indian Railway website www.ireps.gov.in."
    ]
    for i, c in enumerate(certs, 1):
        elements.append(Paragraph(f"{i}. {c}", value_style))
    
    elements.append(Spacer(1, 20))
    elements.append(Paragraph(f"Digitally Signed By: AMM/{data.get('office', 'Railway Authority')}", label_style))

    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_generic_pdf_buffer(data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('TitleStyle', fontName=REGIONAL_FONT, fontSize=18, alignment=1, spaceAfter=2)
    subtitle_style = ParagraphStyle('SubtitleStyle', fontName=REGIONAL_FONT, fontSize=10, alignment=1, textColor=colors.grey, spaceAfter=20)
    header_style = ParagraphStyle('HeaderStyle', fontName=REGIONAL_FONT, fontSize=11, backgroundColor=colors.lightgrey, textColor=colors.black, borderPadding=5, spaceBefore=8, spaceAfter=5)
    cell_style = ParagraphStyle('CellStyle', fontName=REGIONAL_FONT, fontSize=9)
    body_style = ParagraphStyle('BodyStyle', fontName=REGIONAL_FONT, fontSize=10, leading=12)
    
    elements = []
    
    logo_path = find_logo_path()
    logo = None
    if logo_path:
        logo = Image(logo_path, width=60, height=60)
        
    portal_label = data.get('source_label', 'Government e-Marketplace')

    title_info = [[logo, [Paragraph("बिड दस्तावेज़ / Bid Document", title_style), Paragraph(portal_label, subtitle_style)]]]
    title_table = Table(title_info, colWidths=[80, 400])
    title_table.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'MIDDLE'), ('ALIGN', (1,0), (1,0), 'LEFT')]))
    elements.append(title_table)
    
    bid_info = [[Paragraph(f"<b>बिड संख्या / BID NUMBER:</b> {data.get('bidNumber')}", cell_style), 
                 Paragraph(f"<b>दिनांक / DATED:</b> {data.get('dated')}", cell_style)]]
    bid_table = Table(bid_info, colWidths=[250, 250])
    bid_table.setStyle(TableStyle([('ALIGN', (0,0), (0,0), 'LEFT'), ('ALIGN', (1,0), (1,0), 'RIGHT')]))
    elements.append(bid_table)
    elements.append(Spacer(1, 10))
    
    section_titles = [
        "1. Tender Basic Details / बिड विवरण",
        "2. Work / Project Details / कार्य विवरण",
        "3. Department & Authority Details / विभाग विवरण",
        "4. Important Tender Dates / महत्वपूर्ण तिथियाँ",
        "5. Financial & Fee Details / वित्तीय विवरण",
        "6. Purchase Preference / खरीद प्राथमिकता",
        "7. Required Tender Documents / आवश्यक दस्तावेज़",
        "8. Bidder Eligibility Criteria / पात्रता मानदंड",
        "9. Bid Control & Auto-Extension / बिड नियंत्रण",
        "10. BOQ / Commercial Details / बीओक्यू विवरण",
        "11. Geographical Details / भौगोलिक विवरण",
        "12. Enquiry / Bid Forms / पूछताछ प्रपत्र",
        "13. Consignee / Delivery Details / परेषिती विवरण",
        "14. Other Important Links / अन्य महत्वपूर्ण लिंक"
    ]
    
    section_data = {
        "1": [["Bid ID", data.get('bidNumber')], ["Category", data.get('category')], ["Bidding Type", data.get('bidType', 'Two Packet')], ["Evaluation Method", data.get('evaluationMethod', 'L1')]],
        "2": [["Project Name", data.get('category')], ["Contract Value", data.get('estimatedValue')], ["Completion Period", data.get('epbgDuration', 'As per NIT')]],
        "3": [["Ministry", data.get('ministry')], ["Department", data.get('department')], ["Organisation", data.get('organisation')], ["Office", data.get('office')], ["Beneficiary", data.get('beneficiary')]],
        "4": [["Start Date", data.get('dated')], ["End Date", data.get('endDate')], ["Opening Date", data.get('openingDate')], ["Validity", data.get('validity')]],
        "5": [["Estimated Value", data.get('estimatedValue')], ["EMD Required", data.get('emdRequired')], ["Advisory Bank", data.get('advisoryBank')], ["ePBG %", data.get('epbgPercentage')]],
        "6": [["MII Preference", data.get('miiPreference')], ["MSE Preference", data.get('msePreference')], ["Price Breakup Required", data.get('priceBreakupRequired')]],
        "7": [["Documents Required", data.get('docRequired', 'Technical Bid Form, Price Bid (BOQ)')]],
        "8": [["Min Avg Annual Turnover", data.get('bidderTurnover')], ["OEM Annual Turnover", data.get('oemTurnover')], ["Years of Experience", data.get('experience')], ["Past Performance %", data.get('pastPerformance')]],
        "9": [["Min Bids for Auto-Extend", data.get('minBidsForAutoExtend')], ["Auto-Extend Days", data.get('autoExtendDays')], ["Max Extension Count", data.get('autoExtensionCount')], ["Tech Clarification Time", data.get('techClarificationTime')]],
        "11": [["Location", data.get('office')]],
        "13": [["Consignee Address", data.get('beneficiary')], ["Delivery Days", "As per NIT"]]
    }

    for idx, title in enumerate(section_titles, 1):
        elements.append(Paragraph(title, header_style))
        if str(idx) == "14":
            elements.append(Paragraph(data.get('clauses', '').replace('\n', '<br/>'), body_style))
        else:
            rows = section_data.get(str(idx), [["Details", "As per NIT"]])
            formatted_rows = [[Paragraph(str(r[0]), cell_style), Paragraph(str(r[1]), cell_style)] for r in rows]
            t = Table(formatted_rows, colWidths=[180, 320])
            t.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('BACKGROUND', (0,0), (0,-1), colors.whitesmoke),
            ]))
            elements.append(t)
        elements.append(Spacer(1, 5))

    elements.append(Spacer(1, 20))
    elements.append(Paragraph("---धन्यवाद / Thank You---", title_style))
    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_word_buffer(data):
    doc = Document()
    portal_label = data.get('source_label', 'Government e-Marketplace')
    
    logo_path = find_logo_path()
    if logo_path:
        doc.add_picture(logo_path, width=Inches(0.8))
    
    doc.add_heading('बिड दस्तावेज़ / Bid Document', 0)
    doc.add_paragraph(portal_label).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run(f"Bid Number: {data.get('bidNumber')}\t\tDated: {data.get('dated')}").bold = True
    
    sections = [
        ("1. Tender Basic Details", [["Bid ID", data.get('bidNumber')], ["Category", data.get('category')], ["Bidding Type", data.get('bidType', 'Two Packet')]]),
        ("2. Work Details", [["Project Name", data.get('category')], ["Estimated Value", data.get('estimatedValue')]]),
        ("3. Department Details", [["Ministry", data.get('ministry')], ["Department", data.get('department')], ["Organisation", data.get('organisation')]]),
        ("8. Terms & Conditions", [["Details", data.get('clauses')]])
    ]

    for title, rows in sections:
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=len(rows), cols=2)
        table.style = 'Table Grid'
        for i, (label, val) in enumerate(rows):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(val)

    doc.add_paragraph("\n---धन्यवाद / Thank You---").alignment = WD_ALIGN_PARAGRAPH.CENTER
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- MASTER PROMPT ---

MASTER_PROMPT = """
You are a Government Tender Intelligence & Document Generator
specialized in Indian e-Procurement Systems:
- Andhra Pradesh eGP (e-Procurement) Portal
- Central Public Procurement Portal (CPPP)
- Government e-Marketplace (GeM)

📥 INPUTS PROVIDED:
1. Tender Source: {{SOURCE_NAME}}
2. Tender Page Data: {{RAW_TENDER_DATA}}

🎯 TASK OBJECTIVES:
1. Extract ALL tender information available from the input.
2. If specific fields are missing in the input but can be inferred from the context (like suggesting a Procurement Category or Technical Parameters based on a Tender Title), use your knowledge as a procurement expert to SUGGEST realistic and compliant values.
3. Normalize and structure the data into a clean, official Tender Summary.
4. For fields that are strictly unknown and cannot be suggested, use "As per NIT".

OUTPUT FORMAT:
Return a JSON object only:
{
    "bidNumber": "...",
    "dated": "...",
    "endDate": "...",
    "openingDate": "...",
    "validity": "...",
    "ministry": "...",
    "department": "...",
    "organisation": "...",
    "office": "...",
    "category": "...",
    "estimatedValue": "...",
    "emdRequired": "...",
    "advisoryBank": "...",
    "epbgPercentage": "...",
    "epbgDuration": "...",
    "beneficiary": "...",
    "miiPreference": "...",
    "msePreference": "...",
    "bidderTurnover": "...",
    "oemTurnover": "...",
    "experience": "...",
    "mseRelaxation": "...",
    "startupRelaxation": "...",
    "pastPerformance": "...",
    "docRequired": "...",
    "techClarificationTime": "...",
    "minBidsForAutoExtend": "...",
    "autoExtendDays": "...",
    "autoExtensionCount": "...",
    "priceBreakupRequired": "...",
    "transactionFee": "...",
    "plCode": "...",
    "approvingAgency": "...",
    "uvamItemId": "...",
    "techWeightage": "...",
    "finWeightage": "...",
    "mandatoryParams": [{"parameter": "...", "requirement": "..."}],
    "consigneeDetails": [{"name": "...", "address": "...", "quantity": "..."}],
    "boqItems": [{"description": "...", "quantity": "...", "unit": "...", "unitPrice": "..."}],
    "enquiryForms": [{"stage": "...", "name": "...", "type": "...", "mandatory": "..."}],
    "ldPercentage": "...",
    "ldMaxCap": "...",
    "inspectionAuthority": "...",
    "inspectionLocation": "...",
    "arbitrationClause": "...",
    "optionClause": "...",
    "clauses": "..."
}
"""

def parse_ai_json(json_str: str):
    try:
        if "```json" in json_str:
            json_str = json_str.split("```json")[-1].split("```")[0].strip()
        elif "```" in json_str:
            json_str = json_str.split("```")[-1].split("```")[0].strip()
        return json.loads(json_str.strip())
    except Exception as e:
        logger.error(f"Failed to parse AI JSON: {e}")
        return None

def get_mock_tender_data(query: str):
    query_lower = query.lower().strip()
    category = "Electronics/General"
    unit = "Nos"
    ministry = "Not Specified"
    params = []
    if any(kw in query_lower for kw in ["refrigerator", "fridge", "freezer", "chiller"]):
        category = "Refrigerators"
        ministry = "Ministry of Health & Family Welfare"
        params = [{"parameter": "Capacity", "requirement": "300L - 400L"}]
    elif any(kw in query_lower for kw in ["furniture", "chair", "table"]):
        category = "Office Furniture"
        ministry = "Ministry of Housing and Urban Affairs"
    elif any(kw in query_lower for kw in ["computer", "laptop", "it"]):
        category = "Computers & IT Accessories"
        ministry = "Ministry of Electronics & IT"
    return {
        "bidNumber": f"GEM/2026/B/{hash(query) % 10000000}",
        "dated": datetime.datetime.now().strftime("%d-%m-%Y"),
        "endDate": (datetime.datetime.now() + datetime.timedelta(days=21)).strftime("%d-%m-%Y"),
        "category": category,
        "title": query.strip().capitalize() if len(query) > 3 else category,
        "estimatedValue": str(len(query) * 10000 if len(query) > 5 else 500000),
        "ministry": ministry,
        "quantity": "100",
        "unit": unit,
        "mandatoryParams": params,
        "clauses": f"Standard GeM GTC applies for {category}.",
        "miiPreference": "Yes",
        "msePreference": "Yes"
    }

@router.post("/analyze-pdf")
async def analyze_tenders_pdf(
    file: UploadFile = File(...),
    source: str = Form("gem")
):
    portal_map = {
        'gem': 'Government e-Marketplace (GeM)',
        'eproc': 'e-Procurement (CPPP/State Portal)',
        'ap_eproc': 'Andhra Pradesh e-Procurement (eGP)',
        'ireps': 'IREPS (Railways)',
        'custom': 'Custom/Manual'
    }
    
    # Save temp file
    import tempfile
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        content = await file.read()
        tmp.write(content)
        temp_path = tmp.name
    
    try:
        pdf_text = extract_pdf_text(temp_path)
        raw_data = f"Uploaded PDF Content:\n{pdf_text[:10000]}"
        prompt = MASTER_PROMPT.replace("{{SOURCE_NAME}}", portal_map.get(source, 'GeM'))
        prompt = prompt.replace("{{RAW_TENDER_DATA}}", raw_data)
        
        extracted = None

        # --- ATTEMPT GROQ ANALYSIS (Prioritized) ---
        if groq_client:
            try:
                chat_completion = groq_client.chat.completions.create(
                    messages=[
                        {"role": "system", "content": "You are a government procurement expert. Return JSON only."},
                        {"role": "user", "content": prompt}
                    ],
                    model="llama-3.3-70b-versatile",
                    temperature=0.1,
                    response_format={"type": "json_object"}
                )
                extracted = parse_ai_json(chat_completion.choices[0].message.content.strip())
            except Exception as e:
                logger.error(f"Groq PDF analysis failed: {e}")

        # --- ATTEMPT CLAUDE ANALYSIS (Fallback 1) ---
        if not extracted and claude_client:
            try:
                message = claude_client.messages.create(
                    model="claude-3-5-sonnet-20241022",
                    max_tokens=3000,
                    system="You are a government procurement expert. Return JSON only.",
                    messages=[{"role": "user", "content": prompt}]
                )
                extracted = parse_ai_json(message.content[0].text.strip())
            except Exception as e:
                logger.error(f"Claude PDF failed: {e}")
        
        # --- ATTEMPT GEMINI ANALYSIS (Fallback 2) ---
        if not extracted and gemini_model:
            try:
                response = gemini_model.generate_content(f"Analyze this tender and return JSON:\n\n{prompt}")
                extracted = parse_ai_json(response.text.strip())
            except Exception as e:
                logger.error(f"Gemini PDF failed: {e}")
        
        if not extracted:
            extracted = get_mock_tender_data("Tender from PDF")
            
        return {"extracted_data": extracted}
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

@router.post("/analyze")
async def analyze_tenders(request_data: dict = Body(...)):
    query = request_data.get("query", "").lower()
    source = request_data.get("source", "gem")
    
    portal_map = {
        'gem': 'Government e-Marketplace (GeM)',
        'eproc': 'e-Procurement (CPPP/State Portal)',
        'ap_eproc': 'Andhra Pradesh e-Procurement (eGP)',
        'ireps': 'IREPS (Railways)',
        'custom': 'Custom/Manual'
    }
    
    pdf_context = ""
    # Try to find a reference PDF in root if needed
    ref_pdf = Path(os.getcwd()) / "GeM-Bidding-8770263.pdf"
    if ref_pdf.exists():
        pdf_context = extract_pdf_text(str(ref_pdf))
    
    raw_data = f"Query: {query}\n\nReference PDF Text:\n{pdf_context[:8000]}"
    prompt = MASTER_PROMPT.replace("{{SOURCE_NAME}}", portal_map.get(source, 'GeM'))
    prompt = prompt.replace("{{RAW_TENDER_DATA}}", raw_data)
    
    extracted = None
    
    # --- ATTEMPT GROQ ANALYSIS (Prioritized) ---
    if groq_client:
        try:
            chat_completion = groq_client.chat.completions.create(
                messages=[
                    {"role": "system", "content": "You are a government procurement expert. Return JSON only."},
                    {"role": "user", "content": prompt}
                ],
                model="llama-3.3-70b-versatile",
                temperature=0.1,
                response_format={"type": "json_object"}
            )
            extracted = parse_ai_json(chat_completion.choices[0].message.content.strip())
        except Exception as e:
            logger.error(f"Groq query analysis failed: {e}")

    # --- ATTEMPT CLAUDE ANALYSIS (Fallback 1) ---
    if not extracted and claude_client:
        try:
            message = claude_client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=3000,
                system="You are a government procurement expert. Return JSON only.",
                messages=[{"role": "user", "content": prompt}]
            )
            clean_json = message.content[0].text.strip()
            extracted = parse_ai_json(clean_json)
        except Exception as e:
            logger.error(f"Claude query analysis failed: {e}")

    # --- ATTEMPT GEMINI ANALYSIS (Fallback 2) ---
    if not extracted and gemini_model:
        try:
            response = gemini_model.generate_content(f"Analyze this tender and return JSON only:\n\n{prompt}")
            clean_json = response.text.strip()
            extracted = parse_ai_json(clean_json)
        except Exception as e:
            logger.error(f"Gemini query analysis failed: {e}")

    if not extracted:
        # Fallback to mock logic if either AI failed or wasn't configured
        extracted = get_mock_tender_data(query)

    # Portal-specific defaults
    defaults = {
        "bidNumber": f"GEM/2026/B/{hash(query) % 10000000}",
        "dated": datetime.datetime.now().strftime("%d-%m-%Y"),
        "endDate": (datetime.datetime.now() + datetime.timedelta(days=15)).strftime("%Y-%m-%d"),
        "openingDate": (datetime.datetime.now() + datetime.timedelta(days=15)).strftime("%Y-%m-%d"),
        "validity": "180",
        "ministry": "Ministry of Home Affairs",
        "department": "Department of Police",
        "organisation": "State Police Headquarters",
        "office": "Andhra Pradesh",
        "category": "Computers & IT Accessories",
        "bidType": "Open",
        "procurementCategory": "Product",
        "bidderTurnover": "25",
        "oemTurnover": "50",
        "experience": "3",
        "mseRelaxation": "Yes",
        "startupRelaxation": "Yes",
        "pastPerformance": "30",
        "estimatedValue": "1500000",
        "emdRequired": "Yes",
        "emdAmount": "30000",
        "emdMode": "Online / BG",
        "emdValidity": "45",
        "advisoryBank": "State Bank of India",
        "epbgPercentage": "3.00",
        "epbgDuration": "12",
        "beneficiary": "Director General of Police",
        "buyerDesignation": "Assistant Commissioner (IT)",
        "buyerAddress": "Police HQ, Amaravati",
        "buyerContact": "0866-2422222",
        "buyerEmail": "it-procurement@appolice.gov.in",
        "buyerGstin": "37AAAAA0000A1Z5",
        "miiPreference": "Yes",
        "msePreference": "Yes",
        "docRequired": "1. OEM Authorization, 2. MSC Certificate, 3. MII Declaration, 4. Past Performance",
        "transactionFee": "0.00",
        "plCode": "11377045",
        "approvingAgency": "DPIIT",
        "itemName": "High-End Laptops",
        "quantity": "50",
        "unit": "Units",
        "installationRequired": "Yes",
        "testingRequired": "Yes",
        "warrantyPeriod": "3",
        "warrantyType": "Comprehensive",
        "deliveryPeriod": "30",
        "evaluationMethod": "L1",
        "reverseAuction": "Enabled",
        "paymentTerms": "100% after delivery and installation",
        "signatoryName": "K. Satyanarayana",
        "signatoryDesignation": "Procurement Head",
        "clauses": "Standard GeM GTC applies. Make in India Class 1 preference applicable. Delivery must be within 30 days of PO.",
        "declarationNoBlacklist": "Yes",
        "declarationTrueInfo": "Yes",
        "declarationNoConflict": "Yes",
        "agreementToTerms": "Yes",
        "techWeightage": "70",
        "finWeightage": "30",
        "mandatoryParams": [
            {"parameter": "Processor", "requirement": "Intel Core i7 or equivalent"},
            {"parameter": "RAM", "requirement": "16GB LPDDR5"},
            {"parameter": "Storage", "requirement": "512GB SSD"},
            {"parameter": "Display", "requirement": "14 inch FHD Anti-glare"}
        ],
        "consigneeDetails": [
            {"name": "General Manager (IT)", "address": "Police Headquarters, Mangalagiri", "quantity": "30"},
            {"name": "Superintendent of Police", "address": "District Office, Vijayawada", "quantity": "20"}
        ],
        "boqItems": [
            {"description": "Laptops with configuration as per specs", "quantity": "50", "unit": "Nos", "unitPrice": "65000"}
        ],
        "ldPercentage": "0.5",
        "ldMaxCap": "10",
        "inspectionAuthority": "Consignee / Buyer",
        "inspectionLocation": "Consignee Site",
        "arbitrationClause": "Arbitration shall be as per Indian Arbitration Act, with jurisdiction in Amaravati, AP.",
        "optionClause": "25% quantity variation allowed as per GeM GTC.",
    }

    # Merge: Extracted data takes priority over defaults, but don't let empty/placeholder values overwrite good defaults
    final_data = {**defaults}
    for k, v in extracted.items():
        if v and str(v).lower() not in ["as per nit", "none", "null", "n/a", ""]:
            final_data[k] = v
            
    final_data.update({
        "source": source,
        "source_label": portal_map.get(source, 'GeM')
    })
    
    return {"status": "success", "extracted_data": final_data}

@router.post("/download")
async def download_document(data: UnifiedDownloadRequest):
    data_dict = data.dict()
    source = data_dict.get('source')
    format_type = data_dict.get('format', 'pdf').lower()
    
    if format_type == "pdf":
        if source == 'ap_eproc':
            buffer = generate_ap_eproc_pdf_buffer(data_dict)
        elif source == 'ireps':
            buffer = generate_ireps_pdf_buffer(data_dict)
        else:
            buffer = generate_generic_pdf_buffer(data_dict)
        mimetype = "application/pdf"
        filename = f"Tender_{data.bidNumber}.pdf"
    else:
        buffer = generate_word_buffer(data_dict)
        mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"Tender_{data.bidNumber}.docx"
        
    return StreamingResponse(
        buffer,
        media_type=mimetype,
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


